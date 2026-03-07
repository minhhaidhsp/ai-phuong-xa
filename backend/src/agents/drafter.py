import os
import re
import json
import uuid
from pathlib import Path
from datetime import datetime

import ollama
from loguru import logger
from docx import Document as DocxDocument

from src.core.config import get_settings

import re as _re
from datetime import datetime as _dt

settings = get_settings()

BASE_DIR     = Path(__file__).resolve().parent.parent.parent
TEMPLATE_DIR = BASE_DIR / "data" / "templates"
OUTPUT_DIR   = BASE_DIR / "data" / "output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

LOAI_VAN_BAN_TU_DO = {
    "thong_bao_bo_sung":    "Thong bao yeu cau bo sung giay to",
    "thong_bao_ket_qua":    "Thong bao ket qua giai quyet ho so",
    "phieu_tiep_nhan":      "Phieu tiep nhan ho so",
    "thu_moi_lam_viec":     "Thu moi lam viec",
    "giay_xac_nhan_cu_tru": "Giay xac nhan cu tru",
}

LLM_DRAFT_PROMPT = """Ban la chuyen vien soan thao van ban hanh chinh UBND phuong/xa TP.HCM.
Loai van ban: {loai}
Thong tin ho so:
{thong_tin}
Yeu cau: Dung the thuc hanh chinh Viet Nam, ngon ngu trang trong, dien day du thong tin."""


def _replace_in_paragraph(paragraph, mapping: dict):
    full_text = "".join(run.text for run in paragraph.runs)
    if not any(k in full_text for k in mapping):
        return
    new_text = full_text
    for k, v in mapping.items():
        new_text = new_text.replace(k, str(v) if v else "")
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_in_table(table, mapping: dict):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _replace_in_paragraph(para, mapping)


def fill_template_docx(template_path: str, mapping: dict, output_filename: str = None) -> str:
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Khong tim thay file mau: {template_path}")
    doc = DocxDocument(template_path)
    for para in doc.paragraphs:
        _replace_in_paragraph(para, mapping)
    for table in doc.tables:
        _replace_in_table(table, mapping)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_paragraph(para, mapping)
        for para in section.footer.paragraphs:
            _replace_in_paragraph(para, mapping)
    if not output_filename:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"vanban_{ts}_{uuid.uuid4().hex[:6]}.docx"
    output_path = OUTPUT_DIR / output_filename
    doc.save(str(output_path))
    logger.info(f"Template filled: {template_path} -> {output_path}")
    return str(output_path)


def get_placeholders_from_template(template_path: str) -> list[str]:
    if not os.path.exists(template_path):
        return []
    doc = DocxDocument(template_path)
    placeholders = set()
    pattern = re.compile(r"\{\{[A-Z0-9_]+\}\}")
    def scan(text): placeholders.update(pattern.findall(text))
    for para in doc.paragraphs:
        scan(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    scan(para.text)
    return sorted(list(placeholders))




async def extract_info_for_template(mo_ta: str, placeholders: list[str]) -> dict:
    """
    Hybrid extract:
    1. Regex trich ngay/thang/nam, so quyet dinh, ho ten chinh xac
    2. LLM chi xu ly cac truong con lai (kho regex hon)
    """
    # ── BUOC 1: Regex trich cac gia tri co pattern ro rang ───────────────
    regex_result = {}

    # Tach tat ca cum ngay/thang/nam dang DD/MM/YYYY hoac DD tháng MM năm YYYY
    dates_found = _re.findall(r'(\d{1,2})[/\-](\d{1,2})[/\-](\d{4})', mo_ta)
    # [('10', '05', '2019'), ('15', '01', '2024'), ('01', '04', '2026'), ('07', '03', '2026')]

    # Gan theo thu tu xuat hien: ngay_sinh, quyet_dinh, cham_dut, ky
    date_keys = [
        ("NGAY_SINH_TRE", "THANG_SINH_TRE", "NAM_SINH_TRE"),
        ("NGAY_QD",       "THANG_QD",       "NAM_QD"),
        ("NGAY_CHAM_DUT", "THANG_CHAM_DUT", "NAM_CHAM_DUT"),
        ("NGAY_KY",       "THANG_KY",       "NAM_KY"),
    ]
    for i, (d, m, y) in enumerate(dates_found):
        if i < len(date_keys):
            nk, tk, yk = date_keys[i]
            regex_result[nk] = d
            regex_result[tk] = m
            regex_result[yk] = y[-2:]  # "2026" -> "26", giu full cho non-NAM_KY

    # So quyet dinh: dang XXX/QD-... hoac so/QD
    qd_match = _re.search(r'(\d+[/\w\-]*QD[/\-\w]*)', mo_ta, _re.IGNORECASE)
    if qd_match:
        regex_result["SO_QUYET_DINH"] = qd_match.group(1)

    # Ho ten: sau "Ba/Ong/Anh/Chi" hoac dau cau
    ten_match = _re.search(r'(?:Ong|Ba|Anh|Chi)\s+([A-Za-zÀ-ỹ\s]{5,40?)(?=[,\s](?:cu tru|muon|sinh|tai))', mo_ta, _re.IGNORECASE)
    if ten_match:
        regex_result["HO_TEN_NGUOI_VIET_DON"] = ten_match.group(1).strip()

    # Ho ten tre: sau "chau/tre em/cho chau"
    tre_match = _re.search(r'(?:chau|cho chau|tre em)\s+([A-Za-zÀ-ỹ\s]{5,40?)(?=[,\s](?:sinh|ngay))', mo_ta, _re.IGNORECASE)
    if tre_match:
        regex_result["HO_TEN_TRE"] = tre_match.group(1).strip()

    # Dia chi: sau "cu tru tai" hoac "tai dia chi"
    dc_match = _re.search(r'(?:cu tru tai|tai)\s+([^,\.]{10,80?)(?=[,\.]|\s(?:muon|voi|can))', mo_ta, _re.IGNORECASE)
    if dc_match:
        regex_result["DIA_CHI_CU_TRU"] = dc_match.group(1).strip()

    logger.info(f"Regex extract: {regex_result}")

    # ── BUOC 2: LLM xu ly cac truong con lai chua co ─────────────────────
    remaining = [p for p in placeholders
                 if p.replace("{{","").replace("}}","") not in regex_result]

    if remaining:
        ph_list = "\n".join([f"- {p}" for p in remaining])
        prompt = f"""Trich xuat thong tin tu mo ta, dien vao cac truong con thieu.

Truong can dien:
{ph_list}

Quy tac:
- KINH_GUI: ten chuc danh + co quan (vi du: "Chu tich UBND Phuong X, Quan Y")
- LY_DO_1/2/3: ly do tuong ung, de trong "" neu khong co
- Chi tra ve JSON thuan tuy, key khong co ngoac nhon

Mo ta: {mo_ta}
JSON:"""
        client = ollama.AsyncClient(host=settings.OLLAMA_BASE_URL)
        response = await client.chat(
            model=settings.LLM_MODEL_FAST,
            messages=[{"role": "user", "content": prompt}],
            options={"temperature": 0.0, "num_ctx": 2048},
        )
        raw = response.message.content.strip()
        if "```json" in raw:
            raw = raw.split("```json")[1].split("```")[0].strip()
        elif "```" in raw:
            raw = raw.split("```")[1].split("```")[0].strip()
        try:
            llm_result = json.loads(raw)
            regex_result.update(llm_result)
        except json.JSONDecodeError:
            logger.warning(f"LLM JSON parse fail: {raw[:200]}")

    # Chuyen sang format {{KEY}}
    return {f"{{{{{k}}}}}": v for k, v in regex_result.items()}

async def draft_tu_do(loai_van_ban: str, thong_tin_ho_so: dict) -> dict:
    ten_loai  = LOAI_VAN_BAN_TU_DO.get(loai_van_ban, loai_van_ban)
    info_text = "\n".join([f"- {k}: {v}" for k, v in thong_tin_ho_so.items() if v])
    prompt    = LLM_DRAFT_PROMPT.format(loai=ten_loai, thong_tin=info_text)
    client    = ollama.AsyncClient(host=settings.OLLAMA_BASE_URL)
    response  = await client.chat(
        model=settings.LLM_MODEL_MAIN,
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user",   "content": f"Soan: {ten_loai}"},
        ],
        options={"temperature": 0.2, "num_ctx": 4096},
    )
    noi_dung = response.message.content.strip()
    logger.info(f"LLM draft '{loai_van_ban}'")
    return {
        "phuong_thuc": "llm_tu_do", "loai_van_ban": loai_van_ban,
        "ten_loai": ten_loai, "noi_dung": noi_dung,
        "file_path": None, "placeholders_da_dien": None,
        "placeholders_chua_dien": None,
        "huong_dan": "Van ban do AI soan, kiem tra truoc khi su dung.",
    }


async def draft_van_ban(
    loai_van_ban: str,
    thong_tin_ho_so: dict,
    template_path: str = None,
    mapping: dict = None,
    mo_ta_tu_do: str = None,
) -> dict:
    # LUONG 1: Co file mau
    if template_path:
        full_path = TEMPLATE_DIR / template_path
        if not full_path.exists():
            raise FileNotFoundError(
                f"Khong tim thay: {template_path}. Copy file vao: {TEMPLATE_DIR}"
            )
        all_ph = get_placeholders_from_template(str(full_path))
        if not mapping and mo_ta_tu_do:
            logger.info("LLM extract thong tin tu mo_ta...")
            mapping = await extract_info_for_template(mo_ta_tu_do, all_ph)
        mapping = mapping or {}
        ts          = datetime.now().strftime("%Y%m%d_%H%M%S")
        ma_hs       = thong_tin_ho_so.get("ma_ho_so", "HS")
        output_name = f"{ma_hs}_{Path(template_path).stem}_{ts}.docx"
        output_path = fill_template_docx(str(full_path), mapping, output_name)
        preview     = DocxDocument(output_path)
        preview_txt = "\n".join([p.text for p in preview.paragraphs if p.text.strip()])
        return {
            "phuong_thuc": "template_docx", "loai_van_ban": loai_van_ban,
            "ten_loai": Path(template_path).stem,
            "noi_dung": preview_txt[:2000],
            "file_path": output_path,
            "placeholders_da_dien": mapping,
            "placeholders_chua_dien": [p for p in all_ph if not mapping.get(p)],
            "huong_dan": "File .docx da tao. Kiem tra truoc khi in.",
        }
    # LUONG 2: LLM tu do
    return await draft_tu_do(loai_van_ban, thong_tin_ho_so)


def get_loai_van_ban() -> dict:
    return LOAI_VAN_BAN_TU_DO


def list_templates() -> list[dict]:
    if not TEMPLATE_DIR.exists():
        return []
    result = []
    for f in TEMPLATE_DIR.glob("*.docx"):
        phs = get_placeholders_from_template(str(f))
        result.append({
            "file_name": f.name, "file_path": f.name,
            "size_kb": round(f.stat().st_size / 1024, 1),
            "placeholders": phs, "so_placeholder": len(phs),
        })
    return result
